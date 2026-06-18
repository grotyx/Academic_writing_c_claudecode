#!/usr/bin/env python3
"""Compile a reviewer response Markdown file to DOCX.

The expected Markdown is intentionally plain so it remains easy to edit:
Reviewer headings, comments, responses, locations, and revised manuscript text
are written as separate paragraphs. The DOCX output follows the local
Author_response_220803_Final.docx style: reviewer comments are regular text;
response, location, revised text, reviewer headings, and closing text are bold.
"""

from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path
from typing import NamedTuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]


class Block(NamedTuple):
    kind: str
    text: str


def normalize_paragraph(text: str) -> str:
    """Remove only structural Markdown wrappers, preserving manuscript text."""
    value = text.strip()
    if value.startswith(">"):
        value = value.lstrip("> ").strip()
    if value.startswith("#"):
        value = value.lstrip("#").strip()
    for marker in ("**", "__", "*", "_"):
        if value.startswith(marker) and value.endswith(marker) and len(value) > len(marker) * 2:
            value = value[len(marker) : -len(marker)].strip()
    return value


def iter_markdown_paragraphs(text: str, keep_change_markers: bool = False) -> list[str]:
    paragraphs: list[str] = []
    buffer: list[str] = []
    in_change_block = False

    def flush() -> None:
        if buffer:
            paragraphs.append(" ".join(buffer).strip())
            buffer.clear()

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if line == "[CHANGE]":
            flush()
            in_change_block = True
            if keep_change_markers:
                paragraphs.append(line)
            continue
        if line == "[/CHANGE]":
            flush()
            if keep_change_markers:
                paragraphs.append(line)
            in_change_block = False
            continue
        if in_change_block and not keep_change_markers:
            continue
        if not line:
            flush()
            continue
        buffer.append(line)

    flush()
    return paragraphs


def parse_response_markdown(text: str, keep_change_markers: bool = False) -> list[Block]:
    blocks: list[Block] = []
    mode = "body"

    for paragraph in iter_markdown_paragraphs(text, keep_change_markers=keep_change_markers):
        clean = normalize_paragraph(paragraph)
        if not clean:
            continue

        revised_match = re.match(r"^revised text:\s*(.*)$", clean, flags=re.IGNORECASE)
        if revised_match:
            revised_text = revised_match.group(1).strip()
            if revised_text:
                blocks.append(Block("revised", revised_text))
            mode = "revised"
            continue
        closing_match = re.match(r"^reviewer closing:\s*(.*)$", clean, flags=re.IGNORECASE)
        if closing_match:
            closing_text = closing_match.group(1).strip()
            if closing_text:
                blocks.append(Block("closing", closing_text))
            mode = "closing"
            continue

        location_match = re.match(r"^location:\s*(.+)$", clean, flags=re.IGNORECASE)
        if location_match:
            blocks.append(Block("location", location_match.group(1).strip()))
            mode = "response"
            continue

        if re.match(r"^reviewer\s*#?\d+\s*:?$", clean, flags=re.IGNORECASE):
            blocks.append(Block("reviewer", clean))
            mode = "body"
            continue

        if re.match(r"^comment\s+[\w.\-)]+", clean, flags=re.IGNORECASE):
            blocks.append(Block("comment", clean))
            mode = "comment"
            continue

        response_match = re.match(r"^response\s*:?\s*(.*)$", clean, flags=re.IGNORECASE)
        if response_match:
            response_text = response_match.group(1).strip()
            text_value = "Response:" if not response_text else f"Response: {response_text}"
            blocks.append(Block("response", text_value))
            mode = "response"
            continue

        if paragraph.lstrip().startswith("#"):
            blocks.append(Block("title", clean))
            mode = "body"
            continue

        if mode == "revised":
            blocks.append(Block("revised", clean))
        elif mode == "closing":
            blocks.append(Block("closing", clean))
        elif mode == "response":
            blocks.append(Block("response", clean))
        else:
            blocks.append(Block("body", clean))

    return blocks


def set_run_font(run, *, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def apply_document_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.0)
        section.right_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)


def add_block(doc: Document, block: Block) -> None:
    bold_kinds = {"title", "reviewer", "response", "location", "revised", "closing"}
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(block.text)
    set_run_font(run, bold=block.kind in bold_kinds)


def compile_docx(input_path: Path, output_path: Path, keep_change_markers: bool = False) -> Path:
    source_text = input_path.read_text(encoding="utf-8")
    blocks = parse_response_markdown(source_text, keep_change_markers=keep_change_markers)

    doc = Document()
    apply_document_style(doc)
    for block in blocks:
        add_block(doc, block)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def default_output_path(input_path: Path, date_suffix: str) -> Path:
    resolved = input_path.resolve()
    try:
        relative = resolved.relative_to((ROOT / "drafts" / "revision").resolve())
        if len(relative.parts) >= 2:
            rev_folder = relative.parts[0]
            return ROOT / "output" / "revision" / rev_folder / f"{resolved.stem}_{date_suffix}.docx"
    except ValueError:
        pass
    return resolved.with_name(f"{resolved.stem}_{date_suffix}.docx")


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Compile a reviewer response Markdown file to DOCX.")
    parser.add_argument("input", type=Path, help="Path to response_letter_REV*.md")
    parser.add_argument("-o", "--output", type=Path, help="Output DOCX path")
    parser.add_argument("--date-suffix", default=datetime.now().strftime("%y%m%d"), help="YYMMDD suffix")
    parser.add_argument(
        "--keep-change-markers",
        action="store_true",
        help="Keep [CHANGE] blocks in the generated DOCX for internal review.",
    )
    return parser


def main() -> int:
    parser = build_arg_parser()
    args = parser.parse_args()
    input_path = args.input
    if not input_path.exists():
        parser.error(f"Input file not found: {input_path}")
    output_path = args.output or default_output_path(input_path, args.date_suffix)
    compile_docx(input_path, output_path, keep_change_markers=args.keep_change_markers)
    print(f"Wrote {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
